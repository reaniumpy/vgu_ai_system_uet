"""Pull text out of pasted input or an uploaded file (.txt / .pdf / .docx).

Hidden-text tricks (white-on-white, tiny fonts) survive extraction: PDF and DOCX
keep the characters in the file regardless of how they're styled, so the guard
still sees the smuggled instructions.
"""

import io


class ExtractionError(Exception):
    """Raised with a plain-language message when a file can't be read."""


def extract_from_upload(filename: str, data: bytes) -> str:
    name = (filename or "").lower()
    if name.endswith(".pdf"):
        return _from_pdf(data)
    if name.endswith(".docx"):
        return _from_docx(data)
    if name.endswith(".txt") or name.endswith(".md") or name.endswith(".text"):
        return _from_txt(data)
    # Unknown/unsupported (e.g. legacy .doc, images) — be explicit, don't guess.
    raise ExtractionError(
        "That file type isn't supported. Please upload a PDF, Word (.docx), or plain-text "
        "(.txt) file, or paste the text directly."
    )


def _from_txt(data: bytes) -> str:
    for encoding in ("utf-8", "utf-16", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


def _from_pdf(data: bytes) -> str:
    from pypdf import PdfReader

    try:
        reader = PdfReader(io.BytesIO(data))
        parts = [page.extract_text() or "" for page in reader.pages]
    except Exception:
        raise ExtractionError(
            "We couldn't read that PDF. It may be scanned (an image) or damaged. Try a "
            "different file, or paste the text directly."
        )
    text = "\n".join(parts).strip()
    if not text:
        raise ExtractionError(
            "That PDF has no readable text — it looks like a scan or image. Please paste the "
            "text directly instead."
        )
    return text


def _from_docx(data: bytes) -> str:
    from docx import Document

    try:
        doc = Document(io.BytesIO(data))
    except Exception:
        raise ExtractionError(
            "We couldn't read that Word file. Make sure it's a .docx (not the older .doc), "
            "or paste the text directly."
        )
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:  # injections love to hide in table cells
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    text = "\n".join(parts).strip()
    if not text:
        raise ExtractionError("That Word file appears to be empty. Please paste the text directly.")
    return text
